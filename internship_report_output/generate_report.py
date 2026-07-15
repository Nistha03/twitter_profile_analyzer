"""
Generate Twitter Profile Analyzer internship technical report.
Based on repository audit - do not run against modified source assumptions.
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(OUTPUT_DIR, "Twitter_Analyzer_Internship_Technical_Report.docx")
AUDIT_PATH = os.path.join(OUTPUT_DIR, "Twitter_Analyzer_Report_Audit.md")
ARCH_PATH = os.path.join(OUTPUT_DIR, "Twitter_Analyzer_Architecture.md")
SCREENSHOT_PATH = os.path.join(OUTPUT_DIR, "Twitter_Analyzer_Screenshot_Checklist.md")

PROJECT_NAME = "Twitter Profile Analyzer"
ORG = "Innefu Labs"
UNIVERSITY = "KIIT University"
PROGRAMME = "B.Tech — Computer Science and Engineering"


def set_cell_shading(cell, color_hex="D9E2F3"):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i])
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            for p in row[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shading)


def add_placeholder_figure(doc, fig_num, title, instructions, caption):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[SCREENSHOT PLACEHOLDER — Figure {fig_num}: {title}]")
    run.bold = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.size = Pt(11)
    box = doc.add_paragraph()
    box.paragraph_format.left_indent = Inches(0.5)
    box.paragraph_format.right_indent = Inches(0.5)
    r = box.add_run(f"Capture instructions: {instructions}")
    r.italic = True
    r.font.size = Pt(10)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"Figure {fig_num}: {caption}")
    cr.italic = True
    cr.font.size = Pt(10)


def setup_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = RGBColor(0, 51, 102)
        if level == 1:
            hs.font.size = Pt(16)
        elif level == 2:
            hs.font.size = Pt(13)
        else:
            hs.font.size = Pt(11)


def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld1 = OxmlElement("w:fldChar")
        fld1.set(qn("w:fldCharType"), "begin")
        run._r.append(fld1)
        run2 = p.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run2._r.append(instr)
        fld2 = OxmlElement("w:fldChar")
        fld2.set(qn("w:fldCharType"), "end")
        run2._r.append(fld2)


def add_header_footer(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    header = section.header
    hp = header.paragraphs[0]
    hp.text = f"{PROJECT_NAME} — Internship Technical Report | {ORG}"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in hp.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(89, 89, 89)


def build_static_toc(doc):
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        ("1.", "Introduction"),
        ("2.", "Project Overview"),
        ("3.", "Problem Statement"),
        ("4.", "Project Objectives"),
        ("5.", "My Role and Work Completed"),
        ("6.", "Technologies Used"),
        ("7.", "System Architecture"),
        ("8.", "Complete Project Workflow"),
        ("9.", "Twitter/X Data Collection"),
        ("10.", "Probing Agent"),
        ("11.", "Profile Discovery"),
        ("12.", "Tweet Collection"),
        ("13.", "Text Preprocessing"),
        ("14.", "N-Gram Dictionary and Matching"),
        ("15.", "User Flagging / Classification Logic"),
        ("16.", "Browser Automation"),
        ("17.", "Session and Authentication Handling"),
        ("18.", "Backend Implementation"),
        ("19.", "Frontend Application"),
        ("20.", "Data Storage and Result Management"),
        ("21.", "Important Project Files"),
        ("22.", "API Endpoints"),
        ("23.", "Challenges I Faced and How I Worked on Them"),
        ("24.", "Commands to Install the Project"),
        ("25.", "Commands to Run the Project"),
        ("26.", "Commands / Steps to Test the Project"),
        ("27.", "Testing Performed"),
        ("28.", "Performance Observations"),
        ("29.", "Screenshots of Working Application"),
        ("30.", "Current Limitations"),
        ("31.", "Future Improvements"),
        ("32.", "Security, Privacy, and Responsible Use"),
        ("33.", "What I Learned During This Project"),
        ("34.", "Conclusion"),
        ("", "Appendix A — Complete Installation Commands"),
        ("", "Appendix B — Complete Run Commands"),
        ("", "Appendix C — Complete Test Steps"),
        ("", "Appendix D — API Endpoint Summary"),
        ("", "Appendix E — Important Files and Functions"),
        ("", "Appendix F — Screenshot Checklist"),
        ("", "Appendix G — Final Submission Checklist"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        if num:
            p.add_run(f"{num} {title}")
        else:
            p.add_run(title)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()


def build_cover_page(doc):
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("INTERNSHIP TECHNICAL WORK REPORT")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()
    fields = [
        ("Project:", PROJECT_NAME),
        ("Submitted by:", "[YOUR FULL NAME — TO BE FILLED]"),
        ("Organization:", ORG),
        ("University:", UNIVERSITY),
        ("Programme:", PROGRAMME),
        ("Internship Role:", "[TO BE FILLED]"),
        ("Internship Duration:", "[TO BE FILLED]"),
        ("Mentor / Manager:", "[TO BE FILLED]"),
        ("Submission Date:", "[TO BE FILLED]"),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = p.add_run(f"{label} ")
        lr.bold = True
        lr.font.size = Pt(12)
        vr = p.add_run(value)
        vr.font.size = Pt(12)

    doc.add_page_break()

    doc.add_heading("Declaration / Report Note", level=1)
    doc.add_paragraph(
        "This report documents the technical work completed during my internship at Innefu Labs. "
        "All descriptions, commands, and architectural details are based on inspection of the actual "
        "project repository as of the submission date. Where information could not be verified from "
        "code or configuration files, it is clearly marked for manual confirmation before final submission."
    )
    doc.add_page_break()


def generate_docx():
    doc = Document()
    setup_styles(doc)
    add_header_footer(doc)
    build_cover_page(doc)
    build_static_toc(doc)

    # 1. INTRODUCTION
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "During my internship at Innefu Labs, I worked on the Twitter Profile Analyzer project — "
        "a Python-based social media intelligence tool that collects public tweet text from Twitter/X "
        "profiles and analyzes it using machine learning and rule-based NLP techniques."
    )
    doc.add_paragraph(
        "The project addresses a practical need in social media monitoring: given a public profile URL, "
        "how can tweet content be collected reliably, stored, classified, and summarized for review? "
        "The repository implements multiple analysis paths: a Streamlit web application and CLI tools "
        "for product taxonomy classification (FastText), semantic metadata extraction (BitNetClient module), "
        "and hate/offensive content detection using TF-IDF n-gram features with a Random Forest classifier "
        "(twitter_ngram_analyzer.py)."
    )
    doc.add_paragraph(
        "This report explains the verified implementation as found in the repository, including installation "
        "and run commands, data flow, limitations, and testing procedures suitable for internship evaluation."
    )

    # 2. PROJECT OVERVIEW
    doc.add_heading("2. Project Overview", level=1)
    doc.add_paragraph(
        f"{PROJECT_NAME} is a Python application for analyzing public Twitter/X profile timelines. "
        "It does not use the official X API; instead, it uses Playwright browser automation to load "
        "profile pages and extract tweet text from the DOM."
    )
    doc.add_heading("Problem Solved", level=2)
    doc.add_paragraph(
        "The system helps analysts and researchers quickly profile a Twitter/X account by collecting "
        "recent tweets and applying automated text analysis — product category labeling, sentiment/topic "
        "metadata, and optional hate/offensive speech flagging."
    )
    doc.add_heading("Primary Inputs and Outputs", level=2)
    add_table(doc,
        ["Input", "Output"],
        [
            ["Public Twitter/X profile URL (x.com or twitter.com)", "Scraped tweet text stored in SQLite"],
            ["Tweet text (from scrape or fallback dataset)", "FastText predicted product taxonomy category"],
            ["Tweet text", "BitNetClient sentiment, topic, and entity labels (rule-based)"],
            ["Tweet text (n-gram analyzer path)", "HOF/NOT classification per tweet with flagged count"],
            ["Analyzed tweet categories", "Matplotlib pie chart saved to outputs/category_distribution.png"],
        ])

    # 3. PROBLEM STATEMENT
    doc.add_heading("3. Problem Statement", level=1)
    challenges = [
        "Collecting tweet text without official API access requires browser automation, which is sensitive to login walls and DOM changes.",
        "Platform layout changes can block scraping, requiring fallback data paths (verified in app.py and web_app.py).",
        "Text must be cleaned differently for FastText taxonomy vs. TF-IDF n-gram classification.",
        "Persistent browser sessions are needed so users can log in manually when required.",
        "Multiple analysis modes exist (taxonomy vs. hate/offensive detection) with different tweet limits and models.",
        "Dependencies used in code (streamlit, pandas, scikit-learn, joblib) are not all listed in requirements.txt.",
    ]
    for c in challenges:
        doc.add_paragraph(c, style="List Bullet")

    # 4. OBJECTIVES
    doc.add_heading("4. Project Objectives", level=1)
    objectives = [
        "Accept a public Twitter/X profile URL and validate its format.",
        "Launch a headed Chromium browser with a persistent user profile for session reuse.",
        "Scrape tweet text from profile timeline DOM elements.",
        "Store profiles and tweets in a local SQLite database.",
        "Classify tweets into Google Product Taxonomy categories using FastText.",
        "Extract sentiment, topic, and named-entity metadata via the BitNetClient module.",
        "Detect hate/offensive themes using TF-IDF n-grams (1–3) and Random Forest on HASOC2021 labels.",
        "Visualize category distribution with a matplotlib pie chart.",
        "Provide both CLI (app.py, twitter_ngram_analyzer.py) and Streamlit (web_app.py) interfaces.",
    ]
    for o in objectives:
        doc.add_paragraph(o, style="List Bullet")

    # 5. MY ROLE
    doc.add_heading("5. My Role and Work Completed", level=1)
    doc.add_paragraph(
        "Contribution details should be reviewed by the student before submission. "
        "Based on repository evidence, the project includes the following work areas:"
    )
    add_table(doc,
        ["Area", "Evidence in Repository"],
        [
            ["Profile URL input and validation", "app.py, web_app.py, tweet_scraper.py validate_url()"],
            ["Browser automation (Playwright)", "fetcher/browser_launcher.py, fetcher/tweet_scraper.py"],
            ["Tweet scraping from profile pages", "TweetScraper.fetch_profile_tweets()"],
            ["SQLite data persistence", "database/database.py — profiles and tweets tables"],
            ["FastText taxonomy classification", "classifier/fasttext_classifier.py"],
            ["N-gram ML flagging pipeline", "twitter_ngram_analyzer.py, classifier/ngram_trainer.py"],
            ["Rule-based semantic analysis", "llm/bitnet_client.py"],
            ["Visualization", "visualization/charts.py"],
            ["Streamlit web UI", "web_app.py"],
            ["CLI analysis entry points", "app.py, twitter_ngram_analyzer.py"],
        ])

    # 6. TECHNOLOGIES
    doc.add_heading("6. Technologies Used", level=1)
    add_table(doc,
        ["Technology", "Purpose", "Where Used (Verified)"],
        [
            ["Python 3", "Core language", "All .py files"],
            ["Playwright 1.42.0", "Browser automation, DOM scraping", "fetcher/browser_launcher.py, tweet_scraper.py"],
            ["Chromium (via Playwright)", "Persistent headed browser context", "BrowserLauncher.get_context()"],
            ["Streamlit", "Web UI framework", "web_app.py (not in requirements.txt — see limitations)"],
            ["FastText (fasttext-wheel 0.9.2)", "Supervised text classification", "classifier/fasttext_classifier.py"],
            ["scikit-learn", "TF-IDF vectorizer, RandomForest, GridSearchCV", "twitter_ngram_analyzer.py, ngram_trainer.py"],
            ["pandas", "HASOC2021 CSV loading", "twitter_ngram_analyzer.py, ngram_trainer.py"],
            ["joblib", "Model persistence (.pkl)", "twitter_ngram_analyzer.py"],
            ["SQLite3", "Local relational storage", "database/database.py"],
            ["matplotlib 3.8.3", "Pie chart generation", "visualization/charts.py"],
            ["numpy 1.26.4", "Numerical dependency", "requirements.txt"],
            ["regex (re)", "URL validation, text cleaning", "tweet_scraper.py, fasttext_classifier.py"],
        ])

  # 7. ARCHITECTURE
    doc.add_heading("7. System Architecture", level=1)
    doc.add_paragraph(
        "The repository implements a modular Python architecture with three entry points sharing common "
        "fetcher, database, classifier, and visualization modules."
    )
    arch_text = """
┌─────────────────────────────────────────────────────────────────┐
│                     ENTRY POINTS                                 │
│  app.py (CLI)  │  web_app.py (Streamlit)  │  twitter_ngram_analyzer.py │
└────────┬────────────────┬──────────────────────────┬────────────┘
         │                │                          │
         ▼                ▼                          ▼
┌────────────────┐ ┌──────────────┐    ┌──────────────────────────┐
│ BrowserLauncher│ │ TweetScraper │    │ TfidfVectorizer +        │
│ (Playwright)   │ │ (DOM extract)│    │ RandomForest (HASOC2021) │
└────────┬───────┘ └──────┬───────┘    └──────────────────────────┘
         │                │
         ▼                ▼
┌────────────────────────────────────────┐
│         DatabaseManager (SQLite)        │
│   profiles table  │  tweets table      │
└────────────────────┬───────────────────┘
                     │
         ┌───────────┼───────────┐
         ▼           ▼           ▼
┌─────────────┐ ┌──────────┐ ┌─────────────┐
│ FastText    │ │ BitNet   │ │ ChartGen    │
│ Classifier  │ │ Client   │ │ (matplotlib)│
└─────────────┘ └──────────┘ └─────────────┘
"""
    add_code_block(doc, arch_text.strip())
    doc.add_paragraph(
        "Note: No probing agent, home-feed crawler, FastAPI backend, or React frontend was found in the repository."
    )

    # 8. WORKFLOW
    doc.add_heading("8. Complete Project Workflow", level=1)
    doc.add_heading("Path A — Streamlit / app.py (Taxonomy + Semantic Analysis)", level=2)
    steps_a = [
        "User provides profile URL.",
        "TweetScraper.validate_url() checks URL pattern against x.com or twitter.com.",
        "BrowserLauncher launches Chromium persistent context (headed, user_data_dir=browser_profile).",
        "page.goto(profile_url) with wait_until=domcontentloaded, then 20-second timeout.",
        "Locator article[data-testid='tweet'] collects up to max_tweets elements.",
        "Tweet text extracted from div[data-testid='tweetText'].",
        "If no tweets scraped, fallback simulation tweets are used (profile-specific in web_app.py).",
        "Profile and raw tweets inserted into SQLite.",
        "Each unprocessed tweet: FastTextClassifier.clean_text() then predict_category().",
        "BitNetClient.process_all_semantic_tasks() adds sentiment, topic, entities.",
        "Results updated in database; ChartGenerator produces pie chart.",
    ]
    for i, s in enumerate(steps_a, 1):
        doc.add_paragraph(f"{i}. {s}")

    doc.add_heading("Path B — twitter_ngram_analyzer.py (Hate/Offensive Detection)", level=2)
    steps_b = [
        "User provides profile URL; scraper fetches up to 25 tweets (default in this file).",
        "Load models/heavy_ngram_model.pkl if exists; otherwise train via GridSearchCV on en_Hasoc2021_train.csv.",
        "For each tweet: RandomForest predicts task_1 label (HOF or NOT).",
        "HOF prediction prints FLAGGED status; displays matched bigrams/trigrams from TF-IDF features.",
        "Summary prints total flagged tweet count.",
    ]
    for i, s in enumerate(steps_b, 1):
        doc.add_paragraph(f"{i}. {s}")

    # 9. DATA COLLECTION
    doc.add_heading("9. Twitter/X Data Collection", level=1)
    add_table(doc,
        ["Aspect", "Verified Implementation"],
        [
            ["Method", "Playwright DOM scraping — no Tweepy, snscrape, or official API"],
            ["Authentication", "Manual login via persistent browser profile (browser_profile/)"],
            ["Browser mode", "Headed (headless=False)"],
            ["Navigation", "page.goto(profile_url, wait_until='domcontentloaded')"],
            ["Wait time", "page.wait_for_timeout(20000) — 20 seconds fixed"],
            ["Tweet selector", "article[data-testid='tweet']"],
            ["Text selector", "div[data-testid='tweetText']"],
            ["Retries", "None explicit; bare except continues/skips failed elements"],
            ["Rate limiting", "Not implemented"],
            ["Fallback", "Hardcoded simulation tweets when scrape returns empty"],
        ])

    # 10. PROBING AGENT
    doc.add_heading("10. Probing Agent", level=1)
    doc.add_paragraph(
        "No probing agent, home-feed crawler, or automated profile-discovery module was found in the repository. "
        "Profiles are provided manually by the user via URL input. A design involving home-feed crawling and "
        "automatic user discovery was discussed in planning but is not implemented in the current codebase."
    )

    # 11. PROFILE DISCOVERY
    doc.add_heading("11. Profile Discovery", level=1)
    doc.add_paragraph(
        "Profile discovery is manual only. The user enters a full profile URL in the CLI prompt or Streamlit "
        "text input. The handle is extracted via profile_url.rstrip('/').split('/')[-1]. "
        "URL validation regex: ^https?://(www\\.)?(twitter\\.com|x\\.com)/[a-zA-Z0-9_]{1,15}/?$"
    )
    doc.add_paragraph(
        "No duplicate-profile queue, verified-account handling, private-account detection, or feed-based "
        "discovery is implemented."
    )

    # 12. TWEET COLLECTION
    doc.add_heading("12. Tweet Collection", level=1)
    add_table(doc,
        ["Parameter", "web_app.py", "app.py", "twitter_ngram_analyzer.py"],
        [
            ["max_tweets default", "5 (comment: quick testing)", "100", "25"],
            ["Replies included", "Not filtered — first N DOM tweets", "Same", "Same"],
            ["Retweets filtered", "No", "No", "No"],
            ["Pinned tweets", "Not specifically handled", "Same", "Same"],
            ["Timestamp captured", "Empty string ''", "Same", "Same"],
            ["Duplicate removal", "No", "No", "No"],
        ])

    # 13. PREPROCESSING
    doc.add_heading("13. Text Preprocessing", level=1)
    doc.add_heading("FastText path (fasttext_classifier.py — clean_text)", level=2)
    doc.add_paragraph("1. Remove URLs (http, www, https patterns)")
    doc.add_paragraph("2. Remove @mentions and #hashtags")
    doc.add_paragraph("3. Collapse whitespace and lowercase")
    doc.add_heading("N-gram ML path", level=2)
    doc.add_paragraph(
        "Raw tweet text is passed directly to TfidfVectorizer without additional cleaning in "
        "twitter_ngram_analyzer.py. The vectorizer uses ngram_range=(1,3) with default sklearn tokenization."
    )

    # 14. N-GRAM
    doc.add_heading("14. N-Gram Dictionary and Matching", level=1)
    doc.add_paragraph(
        "Important clarification: The hate/offensive detection path does NOT use a static n-gram dictionary file. "
        "It uses machine-learned TF-IDF n-gram features trained on en_Hasoc2021_train.csv (HASOC 2021 dataset). "
        "Labels: task_1 column — HOF (Hate/Offensive) and NOT (Not offensive)."
    )
    doc.add_heading("N-gram Types", level=2)
    doc.add_paragraph("Unigrams, bigrams, and trigrams: ngram_range=(1,3) in twitter_ngram_analyzer.py")
    doc.add_paragraph("ngram_trainer.py uses ngram_range=(2,3) only (bigrams and trigrams)")
    doc.add_heading("Example (Based on Implementation Logic)", level=2)
    doc.add_paragraph("Tweet: \"The weather is absolutely amazing and clear today.\"")
    doc.add_paragraph("The TfidfVectorizer tokenizes and generates features such as unigrams (weather, amazing), bigrams (absolutely amazing), and trigrams from the training vocabulary.")
    doc.add_paragraph("The trained RandomForest predicts NOT → SAFE.")
    doc.add_paragraph("Tweet: \"Anti national elements are spreading fake news everywhere.\"")
    doc.add_paragraph("Model may predict HOF → FLAGGED (Hate/Offensive Theme). Matching n-gram features are printed from the sparse vector (up to 4 bigrams and 4 trigrams shown).")
    doc.add_paragraph(
        "FastText path (separate from flagging): Uses wordNgrams=1 during training on models/taxonomy_train.txt "
        "(Google Product Taxonomy-style multi-label lines). Predicts product categories, not HOF/NOT."
    )

    # 15. FLAGGING
    doc.add_heading("15. User Flagging / Classification Logic", level=1)
    doc.add_paragraph(
        "Flagging exists only in twitter_ngram_analyzer.py and classifier/ngram_trainer.py."
    )
    add_table(doc,
        ["Rule", "Implementation"],
        [
            ["Flag condition", "RandomForest prediction == 'HOF'"],
            ["Per-tweet decision", "Yes — each tweet classified independently"],
            ["User-level aggregation", "No — only summary count of flagged tweets printed"],
            ["Threshold", "None — hard class label from classifier"],
            ["Confidence score", "Not output in terminal analysis"],
            ["FastText flagging", "No — FastText assigns taxonomy categories only"],
        ])

    # 16. BROWSER AUTOMATION
    doc.add_heading("16. Browser Automation", level=1)
    add_table(doc,
        ["Setting", "Value (from browser_launcher.py)"],
        [
            ["Framework", "Playwright sync_api"],
            ["Launch method", "chromium.launch_persistent_context"],
            ["User data directory", "browser_profile (default)"],
            ["Headless", "False (visible browser window)"],
            ["Extra args", "--disable-blink-features=AutomationControlled"],
            ["Context cleanup", "context.close() in finally block"],
        ])
    doc.add_paragraph(
        "CloakBrowser, Selenium, and fingerprint-spoofing libraries are not used. "
        "Platform detection and login-wall issues are acknowledged in code comments as "
        "'platform layout wall' with fallback datasets."
    )

    # 17. SESSION
    doc.add_heading("17. Session and Authentication Handling", level=1)
    doc.add_paragraph(
        "Authentication relies on a persistent Chromium user data directory (browser_profile/). "
        "When the user runs the scraper, a visible browser opens. If a login page appears, the user "
        "must log in manually. web_app.py provides a configurable wait slider (0–120 seconds, default 90) "
        "to allow time for login and tweet loading before scraping proceeds."
    )
    doc.add_paragraph(
        "No cookies, tokens, or credentials are stored in source code. Session state lives in the "
        "browser profile directory and must not be included in submission ZIP files."
    )

    # 18. BACKEND
    doc.add_heading("18. Backend Implementation", level=1)
    doc.add_paragraph(
        "There is no separate FastAPI, Flask, or Django backend. The Streamlit application (web_app.py) "
        "runs analysis logic directly in the Streamlit process. Database access is via SQLite through "
        "DatabaseManager. No REST API routes or WebSocket endpoints exist."
    )

    # 19. FRONTEND
    doc.add_heading("19. Frontend Application", level=1)
    doc.add_paragraph("Framework: Streamlit (web_app.py)")
    add_table(doc,
        ["UI Element", "Function"],
        [
            ["Page title", "Twitter Profile Analyzer"],
            ["Sidebar", "Setup instructions, tweet fetching steps, BitNet note"],
            ["Profile URL input", "st.text_input with x.com/twitter.com placeholder"],
            ["Wait time slider", "0–120 seconds for login/load wait"],
            ["Button 1", "Fetch Tweets — launches scraper, stores in DB"],
            ["Button 2", "Analyze & Chart — runs FastText + BitNetClient, shows chart and dataframe"],
            ["Session state", "current_profile_id stored between button clicks"],
            ["Error states", "Invalid URL, fetch-before-analyze warning, scrape fallback warning"],
        ])
    doc.add_paragraph("Default Streamlit port: 8501 (Streamlit default — requires manual verification before submission).")

    # 20. STORAGE
    doc.add_heading("20. Data Storage and Result Management", level=1)
    add_table(doc,
        ["Storage", "Location", "Format", "Contents"],
        [
            ["SQLite DB", "database/twitter_analyzer.db", "Relational", "profiles, tweets with analysis fields"],
            ["FastText model", "models/google_taxonomy_model.bin", "Binary", "Generated on first run if missing"],
            ["N-gram ML model", "models/heavy_ngram_model.pkl", "joblib pickle", "Generated after GridSearchCV training"],
            ["Training data", "models/taxonomy_train.txt", "FastText format", "Google taxonomy labels + text"],
            ["HASOC dataset", "en_Hasoc2021_train.csv", "CSV", "text, task_1 (HOF/NOT), task_2"],
            ["Chart output", "outputs/category_distribution.png", "PNG", "Category pie chart"],
            ["Browser session", "browser_profile/", "Chromium profile", "Cookies, local storage — SENSITIVE"],
        ])

    # 21. FILES
    doc.add_heading("21. Important Project Files", level=1)
    add_table(doc,
        ["File", "Purpose", "Key Functions/Classes", "Role"],
        [
            ["app.py", "CLI main pipeline", "main()", "Entry: scrape 100 tweets, analyze, chart"],
            ["web_app.py", "Streamlit UI", "main()", "Entry: web-based fetch and analyze"],
            ["twitter_ngram_analyzer.py", "N-gram flagging CLI", "main()", "Entry: scrape 25, train/load ML, flag"],
            ["fetcher/browser_launcher.py", "Browser setup", "BrowserLauncher.get_context()", "Playwright context"],
            ["fetcher/tweet_scraper.py", "Tweet extraction", "TweetScraper.fetch_profile_tweets()", "DOM scraping"],
            ["database/database.py", "Persistence", "DatabaseManager", "SQLite CRUD"],
            ["classifier/fasttext_classifier.py", "Taxonomy ML", "FastTextClassifier", "Train/load/predict"],
            ["classifier/ngram_trainer.py", "N-gram trainer", "HeavyNgramTrainer", "GridSearch training"],
            ["llm/bitnet_client.py", "Semantic metadata", "BitNetClient", "Rule-based sentiment/topic/entities"],
            ["visualization/charts.py", "Charts", "ChartGenerator", "Pie chart generation"],
        ])

    # 22. API
    doc.add_heading("22. API Endpoints", level=1)
    doc.add_paragraph("No HTTP API endpoints are implemented. The application uses direct Python function calls and Streamlit UI events.")
    add_table(doc,
        ["Method", "Endpoint", "Purpose", "Input", "Output", "Source"],
        [
            ["N/A", "N/A", "No REST API", "N/A", "N/A", "Verified: no FastAPI/Flask routes"],
        ])

    # 23. CHALLENGES
    doc.add_heading("23. Challenges I Faced and How I Worked on Them", level=1)
    add_table(doc,
        ["Challenge", "What I Observed", "Likely/Verified Cause", "What I Tried or Changed", "Current Status"],
        [
            ["Empty scrape results", "No tweets returned from profile page", "Login wall or DOM timing (20s may be insufficient)", "Added fallback simulation datasets; web_app wait slider (0–120s)", "Fallback active; live scrape unreliable without login"],
            ["Platform layout wall", "Warning logged in app.py", "X/Twitter UI changes or auth requirement", "Profile-specific and generic mock tweet data", "Documented in code comments"],
            ["Incomplete requirements.txt", "Import errors for streamlit, sklearn, pandas", "requirements.txt lists only 4 packages", "Manual pip install of missing packages needed", "Requires manual verification before submission"],
            ["Missing setup scripts", "web_app sidebar references scripts/setup_bitnet.ps1", "Scripts folder not present in repository", "Manual install/run commands documented in this report", "Scripts not found — use commands in Section 24–25"],
            ["BitNet naming vs. implementation", "Module named BitNetClient", "Actual implementation is keyword-based heuristics", "Used as lightweight semantic placeholder", "Not a real BitNet LLM integration"],
            ["Long model training", "GridSearchCV 405 fits mentioned", "Heavy hyperparameter grid on HASOC2021", "Model saved to heavy_ngram_model.pkl for reuse", "First run slow; subsequent runs load pickle"],
            ["Browser profile sensitivity", "browser_profile/ contains session data", "Persistent context stores cookies", "Exclude from ZIP submission", "Ongoing — review before submission"],
        ])

    # 24. INSTALL
    doc.add_heading("24. Commands to Install the Project", level=1)
    doc.add_paragraph("Environment: Windows PowerShell. Project root: C:\\twitter_profile_analyzer (1)")
    install_steps = [
        ("1", "1", "C:\\", 'cd "C:\\twitter_profile_analyzer (1)"', "Navigate to project root", "Current directory set to project"),
        ("2", "1", r"C:\twitter_profile_analyzer (1)", "python -m venv venv", "Create virtual environment", "venv folder created"),
        ("3", "1", r"C:\twitter_profile_analyzer (1)", r".\venv\Scripts\Activate.ps1", "Activate virtual environment", "Prompt shows (venv)"),
        ("4", "1", r"C:\twitter_profile_analyzer (1)", "python -m pip install --upgrade pip", "Upgrade pip", "pip upgraded"),
        ("5", "1", r"C:\twitter_profile_analyzer (1)", "pip install -r requirements.txt", "Install listed dependencies", "playwright, fasttext-wheel, matplotlib, numpy installed"),
        ("6", "1", r"C:\twitter_profile_analyzer (1)", "pip install streamlit pandas scikit-learn joblib", "Install additional runtime dependencies", "All imports resolve"),
        ("7", "1", r"C:\twitter_profile_analyzer (1)", "playwright install chromium", "Install Playwright Chromium browser", "Chromium browser binaries downloaded"),
    ]
    add_table(doc, ["Step", "Terminal", "Working Directory", "Command", "Purpose", "Expected Result"], install_steps)

    # 25. RUN
    doc.add_heading("25. Commands to Run the Project", level=1)
    doc.add_heading("Option A — Streamlit Web Application (Recommended UI)", level=2)
    run_streamlit = [
        ("1", "1", r"C:\twitter_profile_analyzer (1)", r".\venv\Scripts\Activate.ps1", "Activate venv", "(venv) active"),
        ("2", "1", r"C:\twitter_profile_analyzer (1)", "streamlit run web_app.py", "Start Streamlit app", "Local URL http://localhost:8501 displayed"),
        ("3", "Browser", "N/A", "Open http://localhost:8501", "Access application", "Twitter Profile Analyzer UI loads"),
        ("4", "Browser", "N/A", "Enter profile URL, click Fetch Tweets", "Scrape tweets", "Browser opens; tweets stored or fallback used"),
        ("5", "Browser", "N/A", "Click Analyze & Chart", "Run analysis", "Pie chart and dataframe displayed"),
    ]
    add_table(doc, ["Step", "Terminal", "Working Directory", "Command", "Purpose", "Expected Result"], run_streamlit)

    doc.add_heading("Option B — CLI Full Pipeline (app.py)", level=2)
    add_code_block(doc, r'cd "C:\twitter_profile_analyzer (1)"' + "\n.\\venv\\Scripts\\Activate.ps1\npython app.py")
    doc.add_paragraph("Enter profile URL when prompted. Fetches up to 100 tweets, runs FastText + BitNetClient, prints summary and generates chart.")

    doc.add_heading("Option C — N-Gram Hate/Offensive Analyzer", level=2)
    add_code_block(doc, r'cd "C:\twitter_profile_analyzer (1)"' + "\n.\\venv\\Scripts\\Activate.ps1\npython twitter_ngram_analyzer.py")
    doc.add_paragraph("Enter profile URL. Fetches up to 25 tweets. Trains or loads heavy_ngram_model.pkl. Prints per-tweet FLAGGED/SAFE status.")

    # 26. TEST
    doc.add_heading("26. Commands / Steps to Test the Project", level=1)
    tests = [
        ("TEST 1", "Virtual environment", "Run .\\venv\\Scripts\\Activate.ps1", "Prompt shows (venv)", "[STUDENT TO FILL AFTER FINAL TEST]"),
        ("TEST 2", "Dependency import", "python -c \"import streamlit, playwright, fasttext, sklearn\"", "No ImportError", "[STUDENT TO FILL AFTER FINAL TEST]"),
        ("TEST 3", "Playwright browser", "playwright install chromium", "Chromium installed", "[STUDENT TO FILL AFTER FINAL TEST]"),
        ("TEST 4", "Streamlit startup", "streamlit run web_app.py", "App at localhost:8501", "[STUDENT TO FILL AFTER FINAL TEST]"),
        ("TEST 5", "URL validation", "Enter invalid URL in UI", "Error: valid Twitter/X URL required", "[STUDENT TO FILL AFTER FINAL TEST]"),
        ("TEST 6", "Tweet fetch", "Enter https://x.com/nasa, Fetch Tweets", "Tweets stored or fallback warning shown", "[STUDENT TO FILL AFTER FINAL TEST]"),
        ("TEST 7", "Analysis pipeline", "Click Analyze & Chart after fetch", "Pie chart and dataframe appear", "[STUDENT TO FILL AFTER FINAL TEST]"),
        ("TEST 8", "FastText model", "First analysis run", "models/google_taxonomy_model.bin created", "[STUDENT TO FILL AFTER FINAL TEST]"),
        ("TEST 9", "N-gram analyzer", "python twitter_ngram_analyzer.py", "Per-tweet HOF/NOT output", "[STUDENT TO FILL AFTER FINAL TEST]"),
        ("TEST 10", "SQLite persistence", "Check database/twitter_analyzer.db after fetch", "profiles and tweets rows exist", "[STUDENT TO FILL AFTER FINAL TEST]"),
        ("TEST 11", "CLI app.py", "python app.py with valid URL", "Console summary and chart path logged", "[STUDENT TO FILL AFTER FINAL TEST]"),
        ("TEST 12", "Missing profile", "Fetch without prior analyze", "Warning: click Fetch Tweets first", "[STUDENT TO FILL AFTER FINAL TEST]"),
    ]
    add_table(doc, ["Test ID", "Objective", "Steps", "Expected Result", "Actual Result"], tests)

    # 27. TESTING TABLE
    doc.add_heading("27. Testing Performed", level=1)
    doc.add_paragraph("Actual results must be filled by the student after running tests on their machine.")
    add_table(doc,
        ["Test ID", "Feature", "Input", "Expected Result", "Observed Result", "Status"],
        [[t[0], t[1], t[2], t[3], t[4], "Pending"] for t in tests]
    )

    # 28. PERFORMANCE
    doc.add_heading("28. Performance Observations", level=1)
    doc.add_paragraph("No formal benchmark was verified in the repository.")
    doc.add_paragraph("Informal timing references found in code:")
    doc.add_paragraph("twitter_ngram_analyzer.py: GridSearchCV training reported in minutes (time.perf_counter equivalent via time.time()).", style="List Bullet")
    doc.add_paragraph("tweet_scraper.py: Fixed 20-second page wait per scrape.", style="List Bullet")
    doc.add_paragraph("fasttext_classifier.py: Training configured with epoch=5, wordNgrams=1 for faster training.", style="List Bullet")
    doc.add_paragraph("ngram_trainer.py comments: 10–20 minute training estimate for GridSearchCV.", style="List Bullet")

    # 29. SCREENSHOTS
    doc.add_heading("29. Screenshots of Working Application", level=1)
    doc.add_paragraph("No application screenshots were found in the repository. Placeholders below must be captured manually.")
    screenshots = [
        (1, "Project Structure", "Open project in VS Code/IDE", "Folder tree showing app.py, web_app.py, fetcher/, classifier/, etc.", "Project repository structure in development environment"),
        (2, "Virtual Environment", "Run .\\venv\\Scripts\\Activate.ps1 in PowerShell", "(venv) prefix visible in terminal", "Python virtual environment activation"),
        (3, "Dependency Installation", "Run pip install commands", "Successful package installation output", "Project dependencies installed"),
        (4, "Streamlit Running", "Run streamlit run web_app.py", "Terminal showing Local URL: http://localhost:8501", "Streamlit development server started"),
        (5, "Main Application UI", "Open browser to localhost:8501", "Twitter Profile Analyzer title, URL input, two buttons", "Main Streamlit application interface"),
        (6, "Browser Session", "Click Fetch Tweets", "Chromium window open on X profile page", "Playwright browser during tweet collection"),
        (7, "Fetch Success", "After Fetch Tweets completes", "Green success message with @handle", "Successful tweet fetch confirmation"),
        (8, "Analysis Complete", "After Analyze & Chart", "Pie chart visible on page", "Tweet category distribution chart"),
        (9, "Analyzed Data Table", "After Analyze & Chart", "Streamlit dataframe with tweet analysis columns", "Analyzed tweet results table"),
        (10, "CLI app.py Output", "Run python app.py", "Console showing tweet summary and chart path", "Command-line analysis output"),
        (11, "N-gram Analyzer Output", "Run python twitter_ngram_analyzer.py", "Terminal showing FLAGGED/SAFE per tweet", "N-gram hate/offensive detection results"),
        (12, "SQLite Database", "Open database/twitter_analyzer.db in DB browser", "profiles and tweets tables with data", "Local database persistence"),
        (13, "Fallback Warning", "Fetch when not logged in", "Yellow warning about platform layout wall", "Scrape fallback behavior"),
        (14, "Error — Invalid URL", "Enter non-Twitter URL", "Red error message in Streamlit", "Input validation error handling"),
        (15, "Chart File Output", "Open outputs/category_distribution.png", "Saved pie chart image", "Generated visualization file"),
    ]
    for fig_num, title, instructions, visible, caption in screenshots:
        add_placeholder_figure(doc, fig_num, title, f"{instructions}. Must be visible: {visible}", caption)

    # 30. LIMITATIONS
    doc.add_heading("30. Current Limitations", level=1)
    limitations = [
        "No probing agent or home-feed-based profile discovery.",
        "No official X/Twitter API integration.",
        "Scraping depends on DOM selectors that may break when X updates its UI.",
        "No retry logic or rate limiting for browser automation.",
        "Fallback simulation data used when live scraping fails — results may not reflect real tweets.",
        "BitNetClient is rule-based, not a real large language model.",
        "requirements.txt is incomplete relative to actual imports.",
        "Referenced PowerShell scripts (scripts/setup_bitnet.ps1, scripts/run_project.ps1) are not in the repository.",
        "No user-level flagging aggregation — only per-tweet in n-gram path.",
        "No formal test suite (pytest/unittest files not found).",
        "Browser profile directory contains sensitive session data.",
        "Training taxonomy_train.txt is very large (~500K+ lines) — first FastText training may be slow.",
        "No Docker, CI/CD, or deployment configuration.",
    ]
    for lim in limitations:
        doc.add_paragraph(lim, style="List Bullet")

    # 31. FUTURE
    doc.add_heading("31. Future Improvements", level=1)
    improvements = [
        "Complete requirements.txt with all dependencies and pinned versions.",
        "Add probing agent for home-feed profile discovery if required by use case.",
        "Replace DOM scraping selectors with more resilient extraction or official API where permitted.",
        "Integrate a real LLM or fine-tuned model for BitNetClient semantic tasks.",
        "Add user-level risk scoring aggregating tweet-level flags.",
        "Implement proper retry, timeout, and rate-limit handling in TweetScraper.",
        "Add automated tests and CI pipeline.",
        "Create missing setup/run PowerShell scripts referenced in web_app.py sidebar.",
        "Add configuration file for max_tweets, wait times, and model paths.",
        "Exclude browser_profile from version control and document session setup.",
    ]
    for imp in improvements:
        doc.add_paragraph(imp, style="List Bullet")

    # 32. SECURITY
    doc.add_heading("32. Security, Privacy, and Responsible Use", level=1)
    doc.add_heading("Implemented Controls", level=2)
    doc.add_paragraph("No credentials hardcoded in source files (verified).", style="List Bullet")
    doc.add_paragraph("Local SQLite storage only — no external data transmission in code.", style="List Bullet")
    doc.add_paragraph("URL validation restricts input to twitter.com/x.com profile patterns.", style="List Bullet")
    doc.add_heading("Recommended Future Controls", level=2)
    doc.add_paragraph(".gitignore for browser_profile/, *.db, models/*.bin, models/*.pkl, outputs/", style="List Bullet")
    doc.add_paragraph("Environment variables for any future API keys.", style="List Bullet")
    doc.add_paragraph("Access control if deployed as a shared web service.", style="List Bullet")
    doc.add_paragraph("Compliance review with X/Twitter Terms of Service before production use.", style="List Bullet")
    doc.add_paragraph("Data minimization — store only required tweet fields.", style="List Bullet")

    # 33. LEARNED
    doc.add_heading("33. What I Learned During This Project", level=1)
    doc.add_paragraph(
        "Through this internship project, I gained hands-on experience building a multi-module Python "
        "application for social media text analysis. I learned how Playwright persistent browser contexts "
        "enable session reuse for platforms that require login, and how DOM-based scraping is fragile "
        "compared to API-based access."
    )
    doc.add_paragraph(
        "On the NLP side, I worked with two different text classification approaches: FastText supervised "
        "learning for product taxonomy labeling, and TF-IDF n-gram features with Random Forest for "
        "hate/offensive speech detection on the HASOC2021 dataset. I understood how n-grams capture "
        "word sequences and how GridSearchCV selects hyperparameters, including the trade-off between "
        "model accuracy and training time."
    )
    doc.add_paragraph(
        "I also integrated SQLite for structured storage, Streamlit for rapid UI prototyping, and matplotlib "
        "for visualization. Debugging browser automation taught me to handle empty scrape results gracefully "
        "with fallback paths and configurable wait times. Overall, the project strengthened my skills in "
        "Python engineering, NLP pipelines, and honest documentation of technical limitations."
    )

    # 34. CONCLUSION
    doc.add_heading("34. Conclusion", level=1)
    doc.add_paragraph(
        f"The {PROJECT_NAME} project is a functional Python-based tool for analyzing public Twitter/X "
        "profile timelines through browser automation, machine learning classification, and data visualization. "
        "It provides Streamlit and CLI interfaces, SQLite persistence, FastText taxonomy labeling, "
        "TF-IDF n-gram hate/offensive detection, and matplotlib charts. While live scraping faces platform "
        "reliability constraints and some components remain placeholders (BitNetClient heuristics, missing "
        "setup scripts), the repository demonstrates a complete internship-scale engineering effort suitable "
        "for further hardening and deployment review."
    )

    # APPENDICES
    doc.add_page_break()
    doc.add_heading("Appendix A — Complete Installation Commands", level=1)
    add_code_block(doc, """cd "C:\\twitter_profile_analyzer (1)"
python -m venv venv
.\\venv\\Scripts\\Activate.ps1
python -m pip install --upgrade pip
pip install -r requirements.txt
pip install streamlit pandas scikit-learn joblib
playwright install chromium""")

    doc.add_heading("Appendix B — Complete Run Commands", level=1)
    add_code_block(doc, """# Streamlit UI
cd "C:\\twitter_profile_analyzer (1)"
.\\venv\\Scripts\\Activate.ps1
streamlit run web_app.py

# CLI taxonomy pipeline
python app.py

# N-gram hate/offensive analyzer
python twitter_ngram_analyzer.py""")

    doc.add_heading("Appendix C — Complete Test Steps", level=1)
    doc.add_paragraph("See Section 26 and Section 27 for full test procedures and results tables.")

    doc.add_heading("Appendix D — API Endpoint Summary", level=1)
    doc.add_paragraph("No REST API endpoints. Application uses Streamlit UI events and CLI input.")

    doc.add_heading("Appendix E — Important Files and Functions", level=1)
    doc.add_paragraph("See Section 21 for the complete file reference table.")

    doc.add_heading("Appendix F — Screenshot Checklist", level=1)
    doc.add_paragraph("See Section 29 and Twitter_Analyzer_Screenshot_Checklist.md for all 15 required captures.")

    doc.add_heading("Appendix G — Final Submission Checklist", level=1)
    checklist = [
        "Word report completed (this document)",
        "Table of Contents visible (static TOC included)",
        "Screenshots inserted (student must capture and replace placeholders)",
        "Install/run/test commands verified on local machine",
        "Project ZIP created excluding venv, node_modules, browser_profile",
        "Secrets removed — no .env with credentials in ZIP",
        ".env reviewed if present",
        "browser_profile/ excluded from ZIP (contains session cookies)",
        "Cache/temp files reviewed",
        "Final ZIP tested after extraction on clean folder",
        "Student placeholders filled: name, role, duration, mentor, date",
    ]
    for item in checklist:
        doc.add_paragraph(item, style="List Bullet")

    add_page_numbers(doc)
    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")


def generate_audit_md():
    content = f"""# Twitter Profile Analyzer — Report Audit

Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

## 1. Repository Files Inspected

### Python Source Files
- app.py
- web_app.py
- twitter_ngram_analyzer.py
- fetcher/browser_launcher.py
- fetcher/tweet_scraper.py
- database/database.py
- classifier/fasttext_classifier.py
- classifier/ngram_trainer.py
- llm/bitnet_client.py
- visualization/charts.py

### Configuration & Data
- requirements.txt
- en_Hasoc2021_train.csv
- models/taxonomy_train.txt (504,196+ lines)

### Browser Profile (Sensitive)
- browser_profile/ (Chromium persistent session data)

### Not Found
- README.md (project root)
- package.json / frontend React app
- scripts/setup_bitnet.ps1
- scripts/run_project.ps1
- Probing agent / crawler modules
- Test files (pytest, unittest)
- Screenshots (*.png, *.jpg)
- google_taxonomy_model.bin
- heavy_ngram_model.pkl
- Docker files

## 2. Important Source Files

| File | Role |
|------|------|
| web_app.py | Streamlit UI entry point |
| app.py | CLI taxonomy pipeline (100 tweets) |
| twitter_ngram_analyzer.py | CLI n-gram HOF/NOT flagging (25 tweets) |
| fetcher/tweet_scraper.py | Playwright DOM scraping |
| fetcher/browser_launcher.py | Persistent Chromium context |
| classifier/fasttext_classifier.py | FastText taxonomy |
| classifier/ngram_trainer.py | GridSearch n-gram trainer |
| database/database.py | SQLite persistence |
| llm/bitnet_client.py | Rule-based semantic heuristics |

## 3. Verified Technologies

| Technology | Evidence |
|------------|----------|
| Python | All .py files |
| Playwright 1.42.0 | requirements.txt, tweet_scraper.py, browser_launcher.py |
| Chromium | launch_persistent_context |
| Streamlit | web_app.py import (NOT in requirements.txt) |
| FastText | fasttext_classifier.py, requirements.txt |
| scikit-learn | twitter_ngram_analyzer.py, ngram_trainer.py |
| pandas | CSV loading |
| joblib | Model pickle save/load |
| SQLite | database.py |
| matplotlib | charts.py |
| numpy | requirements.txt |

### NOT Used (Verified Absent)
- FastAPI, Flask, Django, Uvicorn
- React, Vite, Node.js frontend
- Selenium, CloakBrowser
- Tweepy, snscrape, official X API
- PostgreSQL, MongoDB, Redis
- WebSockets
- spaCy, NLTK, Hugging Face Transformers
- Static n-gram dictionary file

## 4. Verified Browser Automation Method

- **Framework:** Playwright sync_api
- **Launch:** chromium.launch_persistent_context
- **user_data_dir:** browser_profile
- **headless:** False
- **args:** --disable-blink-features=AutomationControlled

## 5. Verified Data Collection Method

- Navigate to user-supplied profile URL
- Wait 20 seconds (wait_for_timeout(20000))
- Select article[data-testid='tweet']
- Extract div[data-testid='tweetText'] inner_text
- No scrolling implemented
- No home feed access

## 6. Verified N-Gram Logic

- **NOT a static dictionary** — ML-based TF-IDF features
- **twitter_ngram_analyzer.py:** TfidfVectorizer(ngram_range=(1,3)) + RandomForestClassifier
- **ngram_trainer.py:** TfidfVectorizer(ngram_range=(2,3)) + RandomForestClassifier
- **Training data:** en_Hasoc2021_train.csv, labels in task_1 (HOF, NOT)
- **Model cache:** models/heavy_ngram_model.pkl

## 7. Verified Flagging Logic

- **Condition:** prediction == "HOF"
- **Scope:** Per-tweet only
- **Output:** Terminal print FLAGGED/SAFE + summary count
- **No user-level flagging**
- **No confidence threshold**
- FastText path does NOT flag users

## 8. Verified Backend Routes

None. No REST API. Streamlit runs in-process.

## 9. Verified Frontend Components

- Streamlit page: Twitter Profile Analyzer
- Sidebar setup instructions
- URL text input
- Wait time slider (0-120s, default 90)
- Fetch Tweets button
- Analyze & Chart button
- Pie chart display (st.image)
- Dataframe display (st.dataframe)

## 10. Verified Commands

```powershell
cd "C:\\twitter_profile_analyzer (1)"
python -m venv venv
.\\venv\\Scripts\\Activate.ps1
pip install -r requirements.txt
pip install streamlit pandas scikit-learn joblib
playwright install chromium
streamlit run web_app.py
python app.py
python twitter_ngram_analyzer.py
```

## 11. Commands Requiring Manual Verification

- Streamlit default port 8501
- Exact Playwright install path on student machine
- First-run FastText training duration on large taxonomy_train.txt
- GridSearchCV training time (10-20 min cited in code)
- Whether live X scraping works with student's browser session

## 12. Screenshots Found

None in repository.

## 13. Screenshots Missing (All 15 — Student Must Capture)

See Twitter_Analyzer_Screenshot_Checklist.md

## 14. Claims Requiring Student Confirmation

- [YOUR FULL NAME]
- Internship Role
- Internship Duration
- Mentor / Manager name
- Submission Date
- Personal contribution ownership per module
- Actual test results (Section 27)
- Whether live scraping succeeded on their machine

## 15. Unresolved Technical Limitations

- No probing agent / home feed crawler
- Incomplete requirements.txt
- Missing scripts/ folder referenced in web_app.py
- BitNetClient is heuristic, not real BitNet LLM
- DOM scraping fragile without login
- No automated tests
- No rate limiting / retries

## 16. Security-Sensitive Files (Exclude from ZIP)

- browser_profile/ (entire directory — cookies, session storage)
- database/twitter_analyzer.db (may contain scraped tweet text)
- venv/ (large, machine-specific)
- __pycache__/ directories
- Any .env file if created with credentials

## 17. Remaining Report Placeholders

- Cover page: student name, role, duration, mentor, date
- Section 27: all test observed results
- Section 29: all 15 screenshot placeholders
- Contribution ownership disclaimer in Section 5
"""
    with open(AUDIT_PATH, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"Saved: {AUDIT_PATH}")


def generate_architecture_md():
    content = """# Twitter Profile Analyzer — Architecture

## Entry Points

1. **web_app.py** — Streamlit UI (5 tweets default)
2. **app.py** — CLI full pipeline (100 tweets)
3. **twitter_ngram_analyzer.py** — CLI hate/offensive detection (25 tweets)

## Data Flow

```
User URL Input
    → TweetScraper.validate_url()
    → BrowserLauncher.get_context() [Playwright Chromium]
    → page.goto(profile_url) + 20s wait
    → DOM extract tweets [max_tweets configurable per entry point]
    → [Fallback simulation if empty]
    → DatabaseManager.insert_profile() + insert_raw_tweets()
    → [Path A] FastTextClassifier + BitNetClient → update_analysis_results()
    → [Path B] TfidfVectorizer + RandomForest → terminal flag output
    → ChartGenerator.generate_category_pie_chart() [Path A only]
```

## Module Dependencies

```
web_app.py / app.py
├── fetcher.browser_launcher.BrowserLauncher
├── fetcher.tweet_scraper.TweetScraper
├── database.database.DatabaseManager
├── classifier.fasttext_classifier.FastTextClassifier
├── llm.bitnet_client.BitNetClient
└── visualization.charts.ChartGenerator

twitter_ngram_analyzer.py
├── fetcher.browser_launcher.BrowserLauncher
├── fetcher.tweet_scraper.TweetScraper
├── pandas, sklearn, joblib
└── en_Hasoc2021_train.csv
```

## Storage

- SQLite: database/twitter_analyzer.db
- Models: models/google_taxonomy_model.bin, models/heavy_ngram_model.pkl
- Charts: outputs/category_distribution.png
- Session: browser_profile/
"""
    with open(ARCH_PATH, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"Saved: {ARCH_PATH}")


def generate_screenshot_checklist():
    items = [
        ("Figure 1", "Project structure in VS Code", r'cd "C:\twitter_profile_analyzer (1)"', "IDE file explorer"),
        ("Figure 2", "venv activation", r".\venv\Scripts\Activate.ps1", "(venv) in PowerShell"),
        ("Figure 3", "pip install success", "pip install -r requirements.txt && pip install streamlit pandas scikit-learn joblib", "Successfully installed messages"),
        ("Figure 4", "Streamlit server", "streamlit run web_app.py", "Local URL: http://localhost:8501"),
        ("Figure 5", "Main UI", "Open http://localhost:8501", "Title, URL input, buttons"),
        ("Figure 6", "Browser during fetch", "Click Fetch Tweets", "Chromium on X profile"),
        ("Figure 7", "Fetch success", "After fetch completes", "Green success banner"),
        ("Figure 8", "Pie chart", "Click Analyze & Chart", "Category distribution chart"),
        ("Figure 9", "Data table", "After analysis", "Streamlit dataframe"),
        ("Figure 10", "CLI app.py", "python app.py", "Console tweet summary"),
        ("Figure 11", "N-gram output", "python twitter_ngram_analyzer.py", "FLAGGED/SAFE lines"),
        ("Figure 12", "SQLite DB", "Open twitter_analyzer.db", "profiles + tweets tables"),
        ("Figure 13", "Fallback warning", "Fetch without login", "Yellow warning message"),
        ("Figure 14", "Invalid URL error", "Enter bad URL", "Red error in UI"),
        ("Figure 15", "Chart PNG file", "Open outputs/category_distribution.png", "Saved pie chart"),
    ]
    lines = ["# Screenshot Checklist\n", "| Figure | Description | Command/Action | Must Show |\n", "|--------|-------------|----------------|----------|\n"]
    for fig, desc, cmd, show in items:
        lines.append(f"| {fig} | {desc} | `{cmd}` | {show} |\n")
    with open(SCREENSHOT_PATH, "w", encoding="utf-8") as f:
        f.writelines(lines)
    print(f"Saved: {SCREENSHOT_PATH}")


if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    generate_docx()
    generate_audit_md()
    generate_architecture_md()
    generate_screenshot_checklist()
    print("All report files generated successfully.")
